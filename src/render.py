"""
تركيب المستند — مقاسات مستخرجة من الـ PDF الأصلي، مش تخمين.

كل رقم تحت اتقاس من Elsayed_Mustafa_FlowCV_Resume.pdf بـ pdfplumber:

    الصفحة        595 × 842 نقطة  (A4)
    الهوامش       39.7 يمين وشمال · 40 فوق
    الاسم         19 عريض · في النص
    المسمى        13 عادي · في النص
    سطر الاتصال   11 · في النص
    عنوان القسم   12 عريض + خط أسود 1.5 نقطة بعرض المحتوى كله
    النص والنقط   11
    سطر الوظيفة   11 عريض + التاريخ على اليمين
    سطر التقنيات  11 مائل
    إزاحة النقطة  9.9 نقطة بعد الهامش
    ارتفاع السطر  14.2 نقطة (ثابت)
    قبل قسم       14.3 · قبل وظيفة 8.3 · بين المهارات 12

مفيش أي توليد هنا. كل كلمة جاية من cv.yaml أو من معرّفات اختارها
الموديل واتفحصت قبل ما توصل.
"""
from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT = "Source Sans 3"        # نفس عيلة SourceSansPro اللي في الـ PDF
INK = RGBColor(0x00, 0x00, 0x00)
RULE = "000000"               # الخط تحت العناوين أسود، مش رمادي

# مقاسات الخط (نقطة)
SZ_NAME, SZ_TITLE, SZ_CONTACT = 19, 13, 11
SZ_HEAD, SZ_BODY = 12, 11

# المسافات (نقطة)
LINE = 14.2                   # ارتفاع السطر الثابت
BEFORE_SECTION = 15.0
AFTER_SECTION = 6.0
BEFORE_ENTRY = 8.3
BETWEEN_SKILLS = 12.0

MARGIN = 39.7 / 72            # 0.551 بوصة
TOP = 43.1 / 72          # أول سطر عند 41.7 في الأصلي
BULLET_INDENT = 9.9 / 72      # 49.6 − 39.7
CONTENT_W = 8.268 - 2 * MARGIN
COL1_W = 268.7 / 72           # عرض عمود المهارات الأول، مقاس من الـ PDF


# ── XML بالإيد ──────────────────────────────────────────────────────────

def _hyperlink(par, text: str, url: str, size=SZ_BODY,
               color=INK, bold=False, italic=False):
    """
    لينك حقيقي: النص بيتعرض والرابط مخفي تحته.
    python-docx مفيهاش API للينكات، فبنضيف علاقة ونلف الـ run بإيدنا.
    """
    rid = par.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        fonts.set(qn(a), FONT)
    rPr.append(fonts)

    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)

    col = OxmlElement("w:color")
    col.set(qn("w:val"), str(color))
    rPr.append(col)

    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    link.append(run)
    par._p.append(link)


def _rule(par) -> None:
    """الخط الأسود تحت عنوان القسم — 1.5 نقطة، بعرض المحتوى كله."""
    pPr = par._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")        # ثُمن نقطة → 12 = 1.5 نقطة
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), RULE)
    bdr.append(bottom)
    pPr.append(bdr)


def _no_borders(table) -> None:
    bdr = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        bdr.append(e)
    table._tbl.tblPr.append(bdr)


# ── أدوات ───────────────────────────────────────────────────────────────

def _p(container, before=0.0, after=0.0, indent=0.0, line=LINE):
    par = container.add_paragraph()
    f = par.paragraph_format
    f.space_before = Pt(before)
    f.space_after = Pt(after)
    if line:
        f.line_spacing = Pt(line)
    if indent:
        f.left_indent = Inches(indent)
    return par


def _run(par, text, size=SZ_BODY, bold=False, italic=False, color=INK):
    r = par.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return r


def _section(doc, title: str) -> None:
    par = _p(doc, before=BEFORE_SECTION, after=AFTER_SECTION)
    _run(par, title, size=SZ_HEAD, bold=True)
    _rule(par)


def _bullet(container, text: str) -> None:
    par = container.add_paragraph(style="List Bullet")
    f = par.paragraph_format
    f.space_before = Pt(0)
    f.space_after = Pt(0)
    f.line_spacing = Pt(LINE)
    # النص عند 49.6 والعلامة قبله بـ 6 نقط — مقاس من الـ PDF:
    # كل أسطر النقطة (الأول والملفوف) بتبدأ من نفس المكان.
    f.left_indent = Inches(BULLET_INDENT)
    f.first_line_indent = Pt(-6)
    _run(par, text)


def _entry(doc, title: str, rest: str, right: str,
           url: str | None = None, first=False):
    """سطر الوظيفة/المشروع: العنوان عريض على الشمال، التاريخ على اليمين."""
    par = _p(doc, before=0 if first else BEFORE_ENTRY, after=0)
    par.paragraph_format.tab_stops.add_tab_stop(
        Inches(CONTENT_W), WD_TAB_ALIGNMENT.RIGHT)
    if url:
        _hyperlink(par, title, url, bold=True)
    else:
        _run(par, title, bold=True)
    if rest:
        _run(par, rest, bold=True)
    if right:
        _run(par, "\t" + right)
    return par


# ── المستند ─────────────────────────────────────────────────────────────

def build(cv: dict, bullets: dict, chosen: list[str],
          headline: str = "", summary: str = "") -> BytesIO:
    doc = Document()

    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(SZ_BODY)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.line_spacing = Pt(LINE)

    for s in doc.sections:
        s.page_width, s.page_height = Inches(8.268), Inches(11.693)   # A4
        s.top_margin = s.bottom_margin = Inches(TOP)
        s.left_margin = s.right_margin = Inches(MARGIN)

    contact = cv.get("contact") or {}
    profile = cv.get("profile") or {}

    # ── الترويسة ──
    # الترويسة مالهاش مسافة سطر ثابتة — 14.2 بتخنق خط 19 و13
    h = _p(doc, after=6.4, line=None)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(h, contact.get("name", ""), size=SZ_NAME, bold=True)

    t = _p(doc, after=2.1, line=None)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(t, headline or profile.get("title", ""), size=SZ_TITLE)

    # سطرين صريحين بنفس تقسيم الأصلي، بدل الاعتماد على اللف التلقائي.
    # اللف بيسيب مسافات الفاصل في آخر السطر فبيبوّظ التوسيط، والتقسيم
    # بيتغير مع أي تعديل في الاسم أو المكان. الصريح أضبط.
    #
    # ارتفاع السطر 24 نقطة — مقاس من الأصلي (السطرين على 91.8 و115.8).
    SEP = "          "

    def centred(after=0.0):
        par = _p(doc, after=after, line=24)
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return par

    row1 = centred()
    n = 0
    if contact.get("email"):
        _hyperlink(row1, contact["email"], f"mailto:{contact['email']}",
                   size=SZ_CONTACT)
        n += 1
    if contact.get("phone"):
        if n:
            _run(row1, SEP, size=SZ_CONTACT)
        _hyperlink(row1, contact["phone"], f"tel:{contact['phone']}",
                   size=SZ_CONTACT)
        n += 1
    if contact.get("location"):
        if n:
            _run(row1, SEP, size=SZ_CONTACT)
        _run(row1, contact["location"], size=SZ_CONTACT)

    row2 = centred(after=3.9)
    n = 0
    if profile.get("notes"):
        _run(row2, profile["notes"], size=SZ_CONTACT)
        n += 1
    for label, url in (contact.get("links") or {}).items():
        if n:
            _run(row2, SEP, size=SZ_CONTACT)
        _hyperlink(row2, label, url, size=SZ_CONTACT)
        n += 1

    # ── Summary ──
    text = summary or profile.get("summary", "")
    if text:
        _section(doc, "Summary")
        _run(_p(doc), re.sub(r"\s+", " ", text).strip())

    # ── ترتيب النقط حسب أهميتها للوظيفة ──
    order = {bid: i for i, bid in enumerate(chosen)}
    picked: dict[str, list[str]] = {}
    for bid in chosen:
        b = bullets.get(bid)
        if b:
            picked.setdefault(b.get("parent_id", ""), []).append(bid)

    # ── Experience ──
    exp = [e for e in (cv.get("experience") or []) if picked.get(e.get("id"))]
    exp.sort(key=lambda e: min(order[b] for b in picked[e["id"]]))
    if exp:
        _section(doc, "Experience")
        for i, e in enumerate(exp):
            _entry(doc, e.get("role", ""),
                   f", {e['company']}" if e.get("company") else "",
                   "  |  ".join(filter(None, [e.get("period"), e.get("mode")])),
                   first=(i == 0))
            for bid in picked[e["id"]]:
                _bullet(doc, bullets[bid]["text"])

    # ── Education ──
    for e in (cv.get("education") or []):
        _section(doc, "Education")
        _entry(doc, e.get("degree", ""),
               f", {e['school']}" if e.get("school") else "",
               e.get("period", ""), first=True)
        head = " ".join(filter(None, [e.get("department"),
                                      "Graduation Project:" if e.get("project") else ""]))
        if e.get("project"):
            _run(_p(doc), f"{head} {e['project']}".strip())
        elif head:
            _run(_p(doc), head)
        break

    # ── Skills — عمودين، زي FlowCV ──
    skills = {k: v for k, v in (cv.get("skills") or {}).items() if v}
    if skills:
        _section(doc, "Skills")
        names = list(skills)
        rows = (len(names) + 1) // 2
        table = doc.add_table(rows=rows, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        _no_borders(table)
        for i, name in enumerate(names):
            r, col = i % rows, i // rows
            cell = table.cell(r, col)
            cell.width = Inches(COL1_W if col == 0 else CONTENT_W - COL1_W)
            head = cell.paragraphs[0]
            f = head.paragraph_format
            f.space_before = Pt(0 if r == 0 else BETWEEN_SKILLS)
            f.space_after = Pt(0)
            f.line_spacing = Pt(LINE)
            _run(head, name.replace("_", " ").title(), bold=True)
            _bullet(cell, ", ".join(str(x) for x in skills[name]))

    # ── Projects ──
    proj = [p for p in (cv.get("projects") or []) if picked.get(p.get("id"))]
    proj.sort(key=lambda p: min(order[b] for b in picked[p["id"]]))
    if proj:
        _section(doc, "Projects")
        for i, pr in enumerate(proj):
            _entry(doc, pr.get("name", ""), "", "",
                   url=pr.get("url"), first=(i == 0))
            if pr.get("stack"):
                _run(_p(doc), pr["stack"], italic=True)
            for bid in picked[pr["id"]]:
                _bullet(doc, bullets[bid]["text"])

    # ── Certificates ──
    certs = cv.get("certificates") or []
    if certs:
        _section(doc, "Certificates")
        for c in certs:
            if isinstance(c, str):
                _bullet(doc, c)
                continue
            par = doc.add_paragraph(style="List Bullet")
            f = par.paragraph_format
            f.space_before = f.space_after = Pt(0)
            f.line_spacing = Pt(LINE)
            f.left_indent = Inches(BULLET_INDENT)
            f.first_line_indent = Pt(-6)
            if c.get("url"):
                _hyperlink(par, c.get("name", ""), c["url"], bold=True)
            else:
                _run(par, c.get("name", ""), bold=True)
            if c.get("issuer"):
                _run(par, f"  —  {c['issuer']}")

    # ── Languages ──
    langs = profile.get("languages") or []
    if langs:
        _section(doc, "Languages")
        _run(_p(doc), "   ".join(str(x) for x in langs))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def filename(job: dict) -> str:
    """اسم ملف نضيف — الشركة والوظيفة، عشان تعرفه من التليجرام."""
    parts = [job.get("company_name") or "", job.get("title") or ""]
    slug = re.sub(r"[^\w\s-]", "", " ".join(parts))
    slug = re.sub(r"\s+", "-", slug.strip())[:60].strip("-")
    return f"Elsayed-Mustafa-{slug or 'CV'}.docx"
